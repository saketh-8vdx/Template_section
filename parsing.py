import requests
import requests
from pathlib import Path
from docx import Document
import json
import streamlit as st

API_KEY= st.secrets["API_KEY"]


def parse_pdf(file_path, api):

    # Step 1 — ask Reducto for a presigned URL
    upload_form = requests.post(
        "https://platform.reducto.ai/upload",
        headers={"Authorization": f"Bearer {api}"},
    ).json()

    print("upload_form : ", upload_form)
    # Step 2 — stream the file to the URL Reducto gave you
    with file_path.open("rb") as f:
        requests.put(upload_form["presigned_url"], data=f)

    # Step 3 — parse (or split / extract) using the file_id

    parse_payload = {
        "options": {
            "ocr_mode": "agentic",
            "extraction_mode": "ocr",
            "chunking": {"chunk_mode": "variable"}, 
            "table_summary": {"enabled": True},
            "figure_summary": { 
                "enabled": True,
                "override": True
            },
            # "filter_blocks": [],
            # "force_url_result": False
        },
        "advanced_options": {
            "ocr_system": "combined", 
            "table_output_format": "json", # "md"
            "merge_tables": True,
            "continue_hierarchy": True,
            "keep_line_breaks": True,
            # "page_range": {},
            # "large_table_chunking": {
            #     "enabled": True,
            #     "size": 50
            # },
            "spreadsheet_table_clustering": "default",
            "add_page_markers": False,
            "remove_text_formatting": False,
            "return_ocr_data": False,
            "filter_line_numbers": False,
            "read_comments": True,
            "persist_results": False,
            "exclude_hidden_sheets": False,
            "exclude_hidden_rows_cols": False,
            "enable_change_tracking": False
        },
        "experimental_options": {
            "enrich": {
                "enabled": True,
                "mode": "standard"
            },
            "native_office_conversion": False,
            "enable_checkboxes": False,
            "enable_equations": False,
            "rotate_pages": True,
            "rotate_figures": False,
            "enable_scripts": False,
            "return_figure_images": False,
            "return_table_images": False,
            "layout_model": "default",
            "embed_text_metadata_pdf": False,
            "danger_filter_wide_boxes": False
        },
        "document_url": upload_form["file_id"],
        "priority": True
    }

    parse = requests.post(
        "https://platform.reducto.ai/parse",
        json=parse_payload,
        headers={"Authorization": f"Bearer {api}"},
    )
    parse_json = parse.json()
    print(parse_json)

    # print("file path : ", FILE_PATH)
    # print("file stem : ", FILE_PATH.stem)

    # # Save output to Word document
    # output_file = Path(f"../../Downloads/ai_analyst_data/parse_imagesoutput_{FILE_PATH.stem}.docx")
    # doc = Document()
    # doc.add_heading('Reducto Parse Output', 0)
    # doc.add_paragraph(f"Original File: {FILE_PATH.name}")
    # doc.add_heading('Full JSON Output', level=1)
    # doc.add_paragraph(json.dumps(parse_json, indent=2))
    # doc.save(output_file)
    return parse_json

def process_json_blocks(output_json):
    """
    Process JSON blocks and organize content by section headers.
    
    Args:
        output_json (dict): JSON object containing result.chunks[].blocks structure
        
    Returns:
        list: List of JSON objects with section headers and their contents
    """
    # Initialize empty list to store all content
    all_content = []
    
    # Get the chunks list from the input JSON
    chunks = output_json.get("result", {}).get("chunks", [])
    
    # Check if chunks is a list
    if not isinstance(chunks, list):
        print("Warning: 'result.chunks' key does not contain a list")
        return all_content
    
    # Iterate through each chunk
    for chunk in chunks:
        # Get the blocks list from the current chunk
        blocks = chunk.get("blocks", [])
        
        # Check if blocks is a list
        if not isinstance(blocks, list):
            continue
        
        current_section = None
        current_contents = []
        
        # Iterate through each JSON object in the blocks list
        for block in blocks:
            # Check if the block is a dictionary and has a "type" key
            if isinstance(block, dict) and "type" in block:
                block_type = block.get("type")
                
                # Check if the type is "Section Header"
                if block_type == "Section Header":
                    # If we have a previous section, add it to all_content
                    if current_section is not None:
                        section_obj = {
                            "section_header_name": current_section,
                            "Contents": current_contents
                        }
                        all_content.append(section_obj)
                    
                    # Get the section name/content
                    section_name = block.get("content", block.get("content", f"Section_{len(all_content) + 1}"))
                    current_section = section_name
                    current_contents = []
                    
                else:
                    # If not a section header, add the block to the current section's content
                    if current_section is not None:
                        content_obj = {
                            "Label": block_type,
                            "Content": block.get("content", block.get("content", ""))
                        }
                        current_contents.append(content_obj)
                    else:
                        # If no section header found yet, create a default section
                        if not all_content or all_content[-1]["section_header_name"] != "Default_Section":
                            section_obj = {
                                "section_header_name": "Default_Section",
                                "Contents": []
                            }
                            all_content.append(section_obj)
                        
                        content_obj = {
                            "Label": block_type,
                            "Content": block.get("content", block.get("content", ""))
                        }
                        all_content[-1]["Contents"].append(content_obj)
        
        # Add the last section from this chunk if it exists
        if current_section is not None and current_contents:
            section_obj = {
                "section_header_name": current_section,
                "Contents": current_contents
            }
            all_content.append(section_obj)
    
    return all_content


