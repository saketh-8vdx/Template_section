import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from parsing import parse_pdf, process_json_blocks
from pathlib import Path

# Dummy placeholder function: replace this with your actual logic
def process_pdf(pdf_path):
    # This is just a dummy structure — replace with your actual output
    return [
        {
            "Section_header": "Executive Summary",
            "Prompts": [
                {
                    "type": "Text",
                    "Prompt": "Summarize the main points of the report.",
                    "Sub-section": "Overview of Key Points"
                },
                {
                    "type": "Figure",
                    "Prompt": "Include a chart showing quarterly growth.",
                    "Sub-section": "Growth Analysis"
                }
            ]
        },
        {
            "Section_header": "Market Analysis",
            "Prompts": [
                {
                    "type": "Text",
                    "Prompt": "Describe the current market trends.",
                    "Sub-section": "Trends Overview"
                }
            ]
        }
    ]

# Function to generate DOCX
def generate_docx(data):
    doc = Document()
    doc.add_heading("Generated Report", 0)

    for section in data:
        section_header = section.get("Section_header", "Untitled Section")
        doc.add_heading(section_header, level=1)

        prompts = section.get("Prompts", [])
        for prompt in prompts:
            type_ = prompt.get("type", "Unknown")
            sub_section = prompt.get("Sub-section", "Untitled")
            prompt_text = prompt.get("Prompt", "")

            doc.add_heading(sub_section, level=2)

            p = doc.add_paragraph()
            run = p.add_run(f"Type: {type_}")
            run.bold = True

            doc.add_paragraph(prompt_text)

    # Save to BytesIO for download
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit UI
st.title("Section Suggestions and Prompt Generations")

uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])

if uploaded_file:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_file_path = tmp_file.name

    # Call your processing function
    output = process_pdf(tmp_file_path)
    API_KEY= st.secrets["API_KEY"]
    file_path = Path(tmp_file_path)
    parsing_output = parse_pdf(file_path,API_KEY)
    parsing_json = process_json_blocks(parsing_output)

    # Generate the DOCX
    docx_file = generate_docx(parsing_json)

    # Provide download button
    st.download_button(
        label="Download DOCX",
        data=docx_file,
        file_name="generated_output.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Optional: Show preview of the data
    st.subheader("Parsed Output")
    st.json(parsing_json)
